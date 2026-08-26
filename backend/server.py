from fastapi import FastAPI, APIRouter, HTTPException, Request, Response, Depends, UploadFile, File, Form
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import io
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import httpx
import requests
import json
import re
from starlette.concurrency import run_in_threadpool
from emergentintegrations.llm.chat import LlmChat, UserMessage

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

app = FastAPI()
api_router = APIRouter(prefix="/api")

EMERGENT_AUTH_URL = "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data"

STATUS_VALUES = ["open", "in_progress", "confirmation_pending", "resolved", "reopened"]
LANE_VALUES = ["RESOLVE", "ACTION", "REVIEW"]
PRIORITY_VALUES = ["P0", "P1", "P2", "P3"]
TEAM_VALUES = ["Maintenance", "Leasing", "Concierge"]

DOCUMENT_TYPES = [
    "Lease", "Resident Handbook", "Pet Policy", "Parking Policy",
    "Amenity Rules", "Maintenance Procedures", "Emergency Procedures", "Move-In/Out Instructions",
]

# ------------------- Object storage (Emergent) -------------------
STORAGE_BASE = (os.environ.get("INTEGRATION_PROXY_URL") or "").strip() or "https://integrations.emergentagent.com"
STORAGE_URL = STORAGE_BASE.rstrip("/") + "/objstore/api/v1/storage"
EMERGENT_KEY = os.environ.get("EMERGENT_LLM_KEY")
APP_NAME = "closeloop"
ALLOWED_EXT = {"pdf", "docx", "txt"}
MIME_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "txt": "text/plain",
}
_storage_key = None


def init_storage(force=False):
    global _storage_key
    if _storage_key and not force:
        return _storage_key
    resp = requests.post(f"{STORAGE_URL}/init", json={"emergent_key": EMERGENT_KEY}, timeout=30)
    resp.raise_for_status()
    _storage_key = resp.json()["storage_key"]
    return _storage_key


def put_object(path, data, content_type):
    key = init_storage()
    resp = requests.put(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key, "Content-Type": content_type},
        data=data, timeout=120,
    )
    if resp.status_code == 404:
        key = init_storage(force=True)
        resp = requests.put(
            f"{STORAGE_URL}/objects/{path}",
            headers={"X-Storage-Key": key, "Content-Type": content_type},
            data=data, timeout=120,
        )
    resp.raise_for_status()
    return resp.json()


def get_object(path):
    key = init_storage()
    resp = requests.get(f"{STORAGE_URL}/objects/{path}", headers={"X-Storage-Key": key}, timeout=60)
    if resp.status_code == 404:
        key = init_storage(force=True)
        resp = requests.get(f"{STORAGE_URL}/objects/{path}", headers={"X-Storage-Key": key}, timeout=60)
    resp.raise_for_status()
    return resp.content, resp.headers.get("Content-Type", "application/octet-stream")


def extract_text(data, ext):
    ext = ext.lower()
    try:
        if ext == "pdf":
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(data))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        if ext == "docx":
            import docx
            d = docx.Document(io.BytesIO(data))
            return "\n".join(p.text for p in d.paragraphs)
        if ext == "txt":
            return data.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"extract_text failed: {e}")
    return ""


def chunk_text(text, size=800, overlap=100):
    text = " ".join(text.split())
    if not text:
        return []
    chunks = []
    i = 0
    while i < len(text):
        chunks.append(text[i:i + size])
        i += size - overlap
    return chunks


def now_iso():
    return datetime.now(timezone.utc).isoformat()


# ------------------- AI intent analysis (Claude Sonnet 4.6) -------------------
AI_CATEGORIES = [
    "maintenance", "complaint", "leasing", "parking", "package", "amenity",
    "visitor", "pet", "noise", "safety", "move-in/out", "policy question", "other",
]
AI_INTENTS = [
    "wants information", "wants something fixed", "wants action taken",
    "wants a status update", "wants permission", "wants a reservation",
    "wants management intervention", "wants to report a problem", "wants to escalate a prior issue",
]

AI_SYSTEM_MESSAGE = (
    "You are an expert triage analyst for a multifamily property management team. "
    "Analyze the resident's message (and any prior context) and respond with ONLY a single JSON "
    "object, no markdown fences and no prose. Use exactly these keys:\n"
    f"- category: one of {AI_CATEGORIES}\n"
    f"- primary_intent: one of {AI_INTENTS}\n"
    "- desired_outcome: a short phrase describing what the resident ultimately wants\n"
    "- urgency: one of [\"emergency\", \"urgent\", \"normal\", \"administrative\"]\n"
    "- priority: one of [\"P0\", \"P1\", \"P2\", \"P3\"] where P0 = emergency with potential immediate "
    "safety risk or serious property damage, P1 = urgent needs rapid attention, P2 = normal routine "
    "operational request, P3 = administrative/information request\n"
    "- sentiment: one of [\"positive\", \"neutral\", \"frustrated\", \"angry\", \"anxious\"]\n"
    "- entities: an array of short strings naming key entities (people, places, objects, dates, amounts)\n"
    "- unit_location: the unit number or location referenced in the message, or \"\" if none\n"
    "- human_judgment_required: boolean, true when the situation needs human judgment "
    "(escalations, repeated/prior complaints, safety or legal matters, ambiguity, strong emotion, "
    "policy exceptions)\n"
    "- human_reason: a short sentence explaining why human judgment is or is not required\n"
    "- safety_risk: boolean, true if there is any physical safety hazard\n"
    "- legal_or_financial: boolean, true if legal action, liability, compensation, refund or money is involved\n"
    "- multiple_residents_affected: boolean, true if the issue likely affects more than one unit/resident\n"
    "- interpersonal: boolean, true if this involves a dispute between people (neighbors, staff)\n"
    "- ai_uncertain: boolean, true if the message is ambiguous or you are unsure how to classify it\n"
    "Detect escalation: if the resident mentions having complained before, repeated attempts, or "
    "prior unresolved issues, set primary_intent to 'wants to escalate a prior issue', raise the "
    "priority, and set human_judgment_required to true."
)


def _parse_ai_json(text):
    if not text:
        return None
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*", "", text).strip()
        text = text.rstrip("`").strip()
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if m:
        text = m.group(0)
    try:
        return json.loads(text)
    except Exception as e:
        logger.error(f"AI JSON parse failed: {e}; raw={text[:300]}")
        return None


async def analyze_message(message, context=None):
    key = os.environ.get("EMERGENT_LLM_KEY")
    if not key:
        return None
    try:
        chat = LlmChat(
            api_key=key,
            session_id=f"triage-{uuid.uuid4().hex[:10]}",
            system_message=AI_SYSTEM_MESSAGE,
        ).with_model("anthropic", "claude-sonnet-4-6")
        prompt = message
        if context:
            prompt = f"Prior messages on this issue (oldest first):\n{context}\n\nNew resident message:\n{message}"
        resp = await chat.send_message(UserMessage(text=prompt))
        return _parse_ai_json(resp)
    except Exception as e:
        logger.error(f"analyze_message failed: {e}")
        return None


def derive_lane(analysis):
    if analysis.get("human_judgment_required"):
        return "REVIEW"
    intent = (analysis.get("primary_intent") or "").lower()
    if intent in (
        "wants something fixed", "wants action taken",
        "wants management intervention", "wants to escalate a prior issue",
    ):
        return "ACTION"
    return "RESOLVE"


async def apply_analysis(issue_id, interaction_id, message, context=None):
    analysis = await analyze_message(message, context)
    if not analysis:
        await db.issues.update_one({"id": issue_id}, {"$set": {"ai_analyzed": False}})
        return None
    priority = analysis.get("priority")
    if priority not in PRIORITY_VALUES:
        priority = "P2"
    entities = analysis.get("entities") or []
    if not isinstance(entities, list):
        entities = [str(entities)]
    updates = {
        "primary_intent": analysis.get("primary_intent"),
        "desired_outcome": analysis.get("desired_outcome"),
        "urgency": analysis.get("urgency"),
        "priority": priority,
        "sentiment": analysis.get("sentiment"),
        "entities": [str(e) for e in entities],
        "human_judgment_required": bool(analysis.get("human_judgment_required")),
        "human_reason": analysis.get("human_reason"),
        "lane": derive_lane(analysis),
        "ai_analyzed": True,
    }
    if analysis.get("category"):
        updates["category"] = analysis["category"]
    if analysis.get("unit_location"):
        updates["ai_location"] = analysis["unit_location"]
    await db.issues.update_one({"id": issue_id}, {"$set": updates})
    if interaction_id:
        await db.interactions.update_one(
            {"id": interaction_id},
            {"$set": {
                "detected_intent": analysis.get("primary_intent"),
                "detected_sentiment": analysis.get("sentiment"),
            }},
        )
    return analysis


# ------------------- Step 4: RESOLVE / ACTION / REVIEW lanes -------------------
_STOPWORDS = set(
    "the a an is are am do does did can could would should i my me you your our of to in on at for "
    "with about how what when where why who this that these those it its and or but if my mine please "
    "hey hi there".split()
)


async def retrieve_chunks(query, top_k=5):
    chunks = await db.document_chunks.find({}, {"_id": 0}).to_list(3000)
    if not chunks:
        return []
    doc_ids = list({c["document_id"] for c in chunks})
    docs = await db.property_documents.find(
        {"id": {"$in": doc_ids}, "is_deleted": False}, {"_id": 0}
    ).to_list(1000)
    name_by_id = {d["id"]: d["name"] for d in docs}
    terms = [w for w in re.findall(r"[a-z]+", query.lower()) if w not in _STOPWORDS and len(w) > 2]
    scored = []
    for c in chunks:
        if c["document_id"] not in name_by_id:
            continue
        text = c["text"].lower()
        score = sum(text.count(t) for t in terms)
        if score > 0:
            scored.append((score, c))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [
        {"text": c["text"], "document_id": c["document_id"], "document_name": name_by_id[c["document_id"]]}
        for _, c in scored[:top_k]
    ]


async def _ai_json(system_message, user_text):
    key = os.environ.get("EMERGENT_LLM_KEY")
    if not key:
        return None
    try:
        chat = LlmChat(
            api_key=key, session_id=f"lane-{uuid.uuid4().hex[:10]}", system_message=system_message
        ).with_model("anthropic", "claude-sonnet-4-6")
        resp = await chat.send_message(UserMessage(text=user_text))
        return _parse_ai_json(resp)
    except Exception as e:
        logger.error(f"_ai_json failed: {e}")
        return None


async def answer_from_docs(question):
    chunks = await retrieve_chunks(question, top_k=6)
    if not chunks:
        return None
    context = "\n\n".join(f"[Document: {c['document_name']}]\n{c['text']}" for c in chunks)
    system = (
        "You are a property assistant. Property documents are authoritative. Answer the resident's "
        "question using ONLY the provided document excerpts. Do NOT use outside knowledge or guess. "
        "If the excerpts do not clearly contain the answer, set supported=false. If two or more documents "
        "give CONTRADICTORY rules for this question, set conflict=true and list the conflicting document names. "
        'Respond with ONLY a JSON object: {"supported": boolean, "answer": string, '
        '"relevant_passage": string (the exact sentence or section you used), "source_document": string, '
        '"confidence": "high"|"medium"|"low", "conflict": boolean, "conflicting_documents": array of strings}.'
    )
    return await _ai_json(system, f"Property document excerpts:\n{context}\n\nResident question: {question}")


async def generate_review_summary(message, analysis, context, policy_excerpt):
    system = (
        "You are a triage assistant preparing a concise brief for a human property manager. "
        "Respond with ONLY a JSON object with these exact keys: what_happened, resident_wants, "
        "relevant_history, relevant_policy, why_human_needed, suggested_next_action. Keep each value to "
        "1-2 sentences. Base relevant_policy ONLY on the provided policy excerpts; if none apply, use "
        "'No specific policy found.'"
    )
    payload = (
        f"Resident message: {message}\n\n"
        f"AI analysis: {json.dumps(analysis)}\n\n"
        f"Conversation history (oldest first):\n{context or 'None'}\n\n"
        f"Policy excerpts:\n{policy_excerpt or 'None'}"
    )
    return await _ai_json(system, payload)


def route_team(category):
    c = (category or "").lower()
    if c == "maintenance":
        return "Maintenance"
    if c in ("leasing", "move-in/out"):
        return "Leasing"
    return "Concierge"  # packages, visitors, amenities, access, front-desk & operational default


def decide_lane(analysis):
    if analysis.get("priority") == "P0" or analysis.get("human_judgment_required"):
        return "REVIEW"
    intent = (analysis.get("primary_intent") or "").lower()
    category = (analysis.get("category") or "").lower()
    if intent in ("wants information", "wants permission") or category == "policy question":
        return "RESOLVE"
    return "ACTION"


async def run_triage(issue_id, interaction_id, message, context=None):
    analysis = await analyze_message(message, context)
    if not analysis:
        await db.issues.update_one({"id": issue_id}, {"$set": {"ai_analyzed": False}})
        return None

    priority = analysis.get("priority") if analysis.get("priority") in PRIORITY_VALUES else "P2"
    entities = analysis.get("entities") or []
    if not isinstance(entities, list):
        entities = [str(entities)]
    updates = {
        "primary_intent": analysis.get("primary_intent"),
        "desired_outcome": analysis.get("desired_outcome"),
        "urgency": analysis.get("urgency"),
        "priority": priority,
        "sentiment": analysis.get("sentiment"),
        "entities": [str(e) for e in entities],
        "human_judgment_required": bool(analysis.get("human_judgment_required")),
        "human_reason": analysis.get("human_reason"),
        "ai_analyzed": True,
        "is_emergency": priority == "P0",
        "safety_risk": bool(analysis.get("safety_risk")),
        "legal_or_financial": bool(analysis.get("legal_or_financial")),
        "multiple_residents_affected": bool(analysis.get("multiple_residents_affected")),
        "interpersonal": bool(analysis.get("interpersonal")),
        "ai_uncertain": bool(analysis.get("ai_uncertain")),
    }
    if analysis.get("category"):
        updates["category"] = analysis["category"]
    updates["ai_location"] = analysis.get("unit_location") or None

    if interaction_id:
        await db.interactions.update_one(
            {"id": interaction_id},
            {"$set": {"detected_intent": analysis.get("primary_intent"),
                      "detected_sentiment": analysis.get("sentiment")}},
        )

    lane = decide_lane({**analysis, "priority": priority})

    # RESOLVE candidate: property documents are authoritative
    if lane == "RESOLVE":
        ans = await answer_from_docs(message)
        if ans and ans.get("conflict"):
            summary = await generate_review_summary(message, analysis, context, "")
            updates.update({
                "lane": "REVIEW", "status": "open", "assigned_team": None,
                "human_judgment_required": True, "policy_conflict": True,
                "conflicting_documents": [str(d) for d in (ans.get("conflicting_documents") or [])],
                "human_reason": "Uploaded property documents conflict; staff must resolve the policy conflict.",
                "review_summary": summary,
            })
            await db.issues.update_one({"id": issue_id}, {"$set": updates})
            await finalize_attention(issue_id)
            return analysis
        if ans and ans.get("supported") and ans.get("confidence") == "high" and ans.get("answer"):
            updates.update({
                "lane": "RESOLVE", "status": "resolved", "resolved_at": now_iso(), "assigned_team": None,
                "auto_response": ans.get("answer"), "answer_source": ans.get("source_document"),
                "answer_passage": ans.get("relevant_passage"), "answer_confidence": "high",
                "human_judgment_required": False,
            })
            await db.issues.update_one({"id": issue_id}, {"$set": updates})
            await db.interactions.insert_one(Interaction(issue_id=issue_id, sender="ai", message=ans.get("answer")).model_dump())
            await finalize_attention(issue_id)
            return analysis
        # Medium/Low confidence or unsupported -> REVIEW (never auto-send)
        lane = "REVIEW"
        updates["human_judgment_required"] = True
        similar = await find_similar_resolutions(analysis.get("category"), message)
        if ans and ans.get("answer") and ans.get("confidence") in ("medium", "low"):
            updates["suggested_response"] = ans.get("answer")
            updates["answer_source"] = ans.get("source_document")
            updates["answer_passage"] = ans.get("relevant_passage")
            updates["answer_confidence"] = ans.get("confidence")
            updates["human_reason"] = f"A draft answer exists but confidence is {ans.get('confidence')}; staff must confirm before sending."
        elif similar:
            updates["suggested_response"] = similar[0].get("previous_answer")
            updates["human_reason"] = "No authoritative document answer; similar past resolutions are available for reference."
        else:
            updates["human_reason"] = "The question could not be confidently answered from the property documents."
        if similar:
            updates["similar_cases"] = similar

    if lane == "ACTION":
        team = route_team(analysis.get("category"))
        ack = f"Thanks — your request has been sent to {team}. We'll keep you updated here."
        updates.update({"lane": "ACTION", "assigned_team": team, "status": "open", "acknowledgement": ack})
        await db.issues.update_one({"id": issue_id}, {"$set": updates})
        await db.interactions.insert_one(Interaction(issue_id=issue_id, sender="system", message=ack).model_dump())
        await finalize_attention(issue_id)
        return analysis

    # REVIEW: staff brief + relevant policy (incl. emergency procedures for P0)
    chunks = await retrieve_chunks(message, top_k=4)
    policy_excerpt = "\n\n".join(f"[{c['document_name']}] {c['text']}" for c in chunks) if chunks else ""
    summary = await generate_review_summary(message, analysis, context, policy_excerpt)
    no_policy = (not summary) or ("no specific policy" in ((summary.get("relevant_policy") or "").lower() if summary else ""))
    updates.update({
        "lane": "REVIEW", "assigned_team": None, "human_judgment_required": True,
        "review_summary": summary,
        "review_policy_source": [] if no_policy else [c["document_name"] for c in chunks][:3],
    })
    if priority == "P0":
        updates["is_emergency"] = True
    if not updates.get("similar_cases"):
        sim = await find_similar_resolutions(analysis.get("category"), message)
        if sim:
            updates["similar_cases"] = sim
    await db.issues.update_one({"id": issue_id}, {"$set": updates})
    await finalize_attention(issue_id)
    return analysis


def bump_priority(p):
    order = ["P3", "P2", "P1", "P0"]
    i = order.index(p) if p in order else 1
    return order[min(i + 1, 2)]  # raise one level, cap at P1 (never auto-declare P0 emergency)


def compute_attention(issue):
    score = {"P0": 60, "P1": 38, "P2": 18, "P3": 6}.get(issue.get("priority") or "P2", 18)
    reasons = []
    if issue.get("priority") == "P0":
        reasons.append("P0 emergency / potential safety risk")
    sent = issue.get("sentiment")
    if sent == "angry":
        score += 18; reasons.append("Resident is angry")
    elif sent == "frustrated":
        score += 12; reasons.append("Resident is frustrated")
    elif sent == "anxious":
        score += 8; reasons.append("Resident is anxious")
    ra = issue.get("resolution_attempts") or 0
    if ra > 0:
        score += min(ra * 14, 42); reasons.append(f"{ra} previous resolution attempt(s) failed")
    if issue.get("failed_resolution"):
        score += 8; reasons.append("Previous resolution failed")
    cc = issue.get("contact_count") or 1
    if cc > 1:
        score += min((cc - 1) * 8, 24); reasons.append(f"{cc} resident contacts on this issue")
    if issue.get("human_judgment_required"):
        score += 8; reasons.append("Human judgment flagged by AI")
    if issue.get("safety_risk"):
        score += 12; reasons.append("Safety risk language detected")
    if issue.get("legal_or_financial"):
        score += 14; reasons.append("Legal / financial / compensation language")
    if issue.get("multiple_residents_affected"):
        score += 8; reasons.append("May affect multiple residents")
    if issue.get("interpersonal"):
        score += 8; reasons.append("Interpersonal dispute")
    if issue.get("ai_uncertain"):
        score += 6; reasons.append("AI uncertainty")
    if issue.get("policy_conflict"):
        score += 30; reasons.append("Policy conflict between documents")
    if issue.get("repeat_complaint"):
        score += 8; reasons.append("Repeat complaint pattern")
    if issue.get("is_emergency"):
        score = max(score, 92)
    if issue.get("policy_conflict"):
        score = max(score, 55)
    score = max(0, min(100, int(round(score))))
    if not reasons:
        reasons.append("Routine request — low attention needed")
    return score, reasons


async def finalize_attention(issue_id):
    issue = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    if not issue:
        return
    score, reasons = compute_attention(issue)
    await db.issues.update_one({"id": issue_id}, {"$set": {"human_attention_score": score, "attention_reasons": reasons}})


async def find_similar_resolutions(category, message, limit=3):
    docs = await db.issues.find({"status": "resolved"}, {"_id": 0}).sort("resolved_at", -1).to_list(300)
    if not docs:
        return []
    terms = [w for w in re.findall(r"[a-z]+", (message or "").lower()) if w not in _STOPWORDS and len(w) > 2]
    scored = []
    for d in docs:
        text = (d.get("description", "") + " " + (d.get("category") or "")).lower()
        score = sum(text.count(t) for t in terms)
        if d.get("category") and category and d["category"] == category:
            score += 3
        if score > 0:
            scored.append((score, d))
    scored.sort(key=lambda x: x[0], reverse=True)
    out = []
    for _, d in scored[:limit]:
        out.append({
            "issue_id": d["id"],
            "issue": d.get("description", "")[:160],
            "previous_answer": d.get("auto_response") or (d.get("review_summary") or {}).get("suggested_next_action") or "Resolved by staff.",
            "outcome": d.get("status"),
            "resident_confirmed": bool(d.get("resident_confirmed")),
        })
    return out


async def match_existing_issue(resident_id, new_message):
    candidates = await db.issues.find({"resident_id": resident_id}, {"_id": 0}).sort("created_at", -1).to_list(10)
    if not candidates:
        return None
    listing = []
    for idx, c in enumerate(candidates):
        listing.append(f"{idx}. [{c['status']}] {c.get('category')} | \"{c['description'][:120]}\" | created {c['created_at'][:10]}")
    system = (
        "You match a new resident message to a prior issue for the SAME resident, to avoid duplicate tickets. "
        'Respond ONLY JSON: {"match_index": integer or -1, "relationship": "same_open"|"returned_resolved"|"none", '
        '"problem_persists": boolean, "reason": string}. '
        "match_index = the number of the existing issue the new message is about, or -1 if it is a genuinely new unrelated topic. "
        "relationship 'same_open' = about an issue still open/in progress; 'returned_resolved' = about a previously resolved issue where the resident now says the problem remains or returned; 'none' = unrelated. "
        "Only match when it is clearly the same underlying problem (same appliance / location / subject)."
    )
    payload = "Existing issues:\n" + "\n".join(listing) + f"\n\nNew resident message: {new_message}"
    res = await _ai_json(system, payload)
    if not res:
        return None
    mi = res.get("match_index", -1)
    if not isinstance(mi, int) or mi < 0 or mi >= len(candidates):
        return None
    return {"issue": candidates[mi], "relationship": res.get("relationship"),
            "persists": bool(res.get("problem_persists")), "reason": res.get("reason")}


async def build_repeat_complaint(issue):
    inters = await db.interactions.find({"issue_id": issue["id"]}, {"_id": 0}).sort("created_at", 1).to_list(500)
    resident_msgs = [i for i in inters if i["sender"] == "resident"]
    actions = [i["message"] for i in inters if i["sender"] in ("staff", "system")]
    # The problem was reported again, so any prior intervention did not hold — derive from history.
    intervention_worked = bool(issue.get("resolved_at")) and bool(issue.get("resident_confirmed"))
    return {
        "first_contact": issue.get("first_reported_at") or issue["created_at"],
        "contact_count": len(resident_msgs),
        "previous_actions": actions[-5:],
        "current_sentiment": issue.get("sentiment"),
        "intervention_worked": intervention_worked,
    }


async def set_review_brief(issue_id, message, context=None):
    issue = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    if not issue:
        return
    chunks = await retrieve_chunks(message, top_k=4)
    policy = "\n\n".join(f"[{c['document_name']}] {c['text']}" for c in chunks) if chunks else ""
    analysis = {
        "category": issue.get("category"), "primary_intent": issue.get("primary_intent"),
        "priority": issue.get("priority"), "sentiment": issue.get("sentiment"),
        "desired_outcome": issue.get("desired_outcome"),
    }
    summary = await generate_review_summary(message, analysis, context, policy)
    no_policy = (not summary) or ("no specific policy" in ((summary.get("relevant_policy") or "").lower() if summary else ""))
    await db.issues.update_one({"id": issue_id}, {"$set": {
        "review_summary": summary,
        "review_policy_source": [] if no_policy else [c["document_name"] for c in chunks][:3],
    }})


# ------------------- Models -------------------
class Property(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    address: str


class Resident(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    unit: str
    property_id: str


class Issue(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    property_id: str
    resident_id: str
    unit: str
    category: Optional[str] = None
    description: str
    desired_outcome: Optional[str] = None
    priority: str = "P2"
    lane: str = "REVIEW"
    assigned_team: Optional[str] = None
    status: str = "open"
    human_reason: Optional[str] = None
    created_at: str = Field(default_factory=now_iso)
    resolved_at: Optional[str] = None
    resident_confirmed: bool = False
    resolution_attempts: int = 0
    # AI intent understanding
    primary_intent: Optional[str] = None
    urgency: Optional[str] = None
    sentiment: Optional[str] = None
    entities: List[str] = Field(default_factory=list)
    ai_location: Optional[str] = None
    human_judgment_required: bool = False
    ai_analyzed: bool = False
    # Lane outcomes (Step 4)
    is_emergency: bool = False
    auto_response: Optional[str] = None
    answer_source: Optional[str] = None
    answer_confidence: Optional[str] = None
    acknowledgement: Optional[str] = None
    review_summary: Optional[Dict[str, Any]] = None
    review_policy_source: List[str] = Field(default_factory=list)
    # Steps 5-12
    answer_passage: Optional[str] = None
    policy_conflict: bool = False
    conflicting_documents: List[str] = Field(default_factory=list)
    suggested_response: Optional[str] = None
    similar_cases: List[Dict[str, Any]] = Field(default_factory=list)
    repeat_complaint: Optional[Dict[str, Any]] = None
    failed_resolution: bool = False
    first_reported_at: Optional[str] = None
    previous_resolved_at: Optional[str] = None
    contact_count: int = 1
    confirmation_requested_at: Optional[str] = None
    confirmed_at: Optional[str] = None
    human_attention_score: int = 0
    attention_reasons: List[str] = Field(default_factory=list)
    safety_risk: bool = False
    legal_or_financial: bool = False
    multiple_residents_affected: bool = False
    interpersonal: bool = False
    ai_uncertain: bool = False
    incident_id: Optional[str] = None


class Interaction(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    issue_id: str
    resident_id: Optional[str] = None
    sender: str  # "resident" | "staff" | "system"
    message: str
    created_at: str = Field(default_factory=now_iso)
    detected_intent: Optional[str] = None
    detected_sentiment: Optional[str] = None


class PropertyDocument(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    property_id: str
    name: str
    doc_type: str
    storage_path: str
    original_filename: str
    content_type: str
    size: int = 0
    processing_status: str = "pending"  # pending | processing | ready | failed
    chunk_count: int = 0
    is_deleted: bool = False
    uploaded_at: str = Field(default_factory=now_iso)
    uploaded_by: Optional[str] = None


# ------------------- Request Schemas -------------------
class IssueCreate(BaseModel):
    name: str
    unit: str
    message: str
    category: Optional[str] = None
    property_id: Optional[str] = None


class IssueUpdate(BaseModel):
    status: Optional[str] = None
    lane: Optional[str] = None
    priority: Optional[str] = None
    assigned_team: Optional[str] = None
    human_reason: Optional[str] = None


class StaffMessage(BaseModel):
    message: str


class ResidentLookup(BaseModel):
    name: str
    unit: str


# ------------------- Auth helpers -------------------
async def get_current_user(request: Request):
    token = request.cookies.get("session_token")
    if not token:
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            token = auth[7:]
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")

    session = await db.user_sessions.find_one({"session_token": token}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=401, detail="Invalid session")

    expires_at = session["expires_at"]
    if isinstance(expires_at, str):
        expires_at = datetime.fromisoformat(expires_at)
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=401, detail="Session expired")

    user = await db.users.find_one({"user_id": session["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User not found")
    return user


# ------------------- Auth routes -------------------
@api_router.post("/auth/session")
async def create_session(request: Request, response: Response):
    session_id = request.headers.get("X-Session-ID")
    if not session_id:
        raise HTTPException(status_code=400, detail="Missing session id")

    async with httpx.AsyncClient() as hc:
        r = await hc.get(EMERGENT_AUTH_URL, headers={"X-Session-ID": session_id})
        if r.status_code != 200:
            raise HTTPException(status_code=401, detail="Invalid session id")
        data = r.json()

    email = data["email"]
    existing = await db.users.find_one({"email": email}, {"_id": 0})
    if existing:
        user_id = existing["user_id"]
        await db.users.update_one(
            {"user_id": user_id},
            {"$set": {"name": data.get("name"), "picture": data.get("picture")}},
        )
    else:
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        await db.users.insert_one({
            "user_id": user_id,
            "email": email,
            "name": data.get("name"),
            "picture": data.get("picture"),
            "created_at": now_iso(),
        })

    session_token = data["session_token"]
    expires_at = datetime.now(timezone.utc) + timedelta(days=7)
    await db.user_sessions.insert_one({
        "user_id": user_id,
        "session_token": session_token,
        "expires_at": expires_at.isoformat(),
        "created_at": now_iso(),
    })

    response.set_cookie(
        key="session_token",
        value=session_token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=7 * 24 * 60 * 60,
    )
    user = await db.users.find_one({"user_id": user_id}, {"_id": 0})
    return user


@api_router.get("/auth/me")
async def auth_me(user=Depends(get_current_user)):
    return user


@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    token = request.cookies.get("session_token")
    if token:
        await db.user_sessions.delete_one({"session_token": token})
    response.delete_cookie("session_token", path="/")
    return {"ok": True}


# ------------------- Public / Resident routes -------------------
@api_router.get("/config")
async def get_config():
    prop = await db.properties.find_one({}, {"_id": 0})
    return {
        "property": prop,
        "statuses": STATUS_VALUES,
        "lanes": LANE_VALUES,
        "priorities": PRIORITY_VALUES,
        "teams": TEAM_VALUES,
        "categories": ["Plumbing", "Electrical", "HVAC", "Appliance", "Pest Control", "General", "Noise", "Other"],
        "document_types": DOCUMENT_TYPES,
    }


@api_router.post("/issues")
async def create_issue(payload: IssueCreate):
    prop = None
    if payload.property_id:
        prop = await db.properties.find_one({"id": payload.property_id}, {"_id": 0})
    if not prop:
        prop = await db.properties.find_one({}, {"_id": 0})
    if not prop:
        raise HTTPException(status_code=400, detail="No property configured")
    property_id = prop["id"]

    name = payload.name.strip()
    unit = payload.unit.strip()
    message = payload.message.strip()
    if not name or not unit or not message:
        raise HTTPException(status_code=400, detail="Name, unit and message are required")

    existing_resident = await db.residents.find_one(
        {"name": name, "unit": unit, "property_id": property_id}, {"_id": 0}
    )
    resident = existing_resident
    if not resident:
        resident = Resident(name=name, unit=unit, property_id=property_id).model_dump()
        await db.residents.insert_one(resident)

    # Resolution Memory: does this relate to an existing issue for this resident?
    match = None
    if existing_resident:
        match = await match_existing_issue(resident["id"], message)

    if match:
        target = match["issue"]
        rel = match.get("relationship")
        interaction = Interaction(issue_id=target["id"], resident_id=resident["id"], sender="resident", message=message)
        await db.interactions.insert_one(interaction.model_dump())
        rmsg_count = await db.interactions.count_documents({"issue_id": target["id"], "sender": "resident"})

        returned = rel == "returned_resolved" or target["status"] in ("resolved", "confirmation_pending")
        if returned and match.get("persists"):
            repeat = await build_repeat_complaint(target)
            await db.issues.update_one({"id": target["id"]}, {"$set": {
                "status": "reopened", "lane": "REVIEW", "resident_confirmed": False,
                "resolution_attempts": (target.get("resolution_attempts") or 0) + 1,
                "failed_resolution": True,
                "previous_resolved_at": target.get("resolved_at"),
                "first_reported_at": target.get("first_reported_at") or target.get("created_at"),
                "contact_count": rmsg_count,
                "priority": bump_priority(target.get("priority") or "P2"),
                "human_judgment_required": True,
                "repeat_complaint": repeat,
                "human_reason": "Resident reports the previously resolved problem has returned or was not fixed.",
            }})
            await db.interactions.insert_one(Interaction(issue_id=target["id"], sender="system", message="Previous resolution failed — issue reopened and routed to REVIEW.").model_dump())
            await set_review_brief(target["id"], message)
        else:
            await db.issues.update_one({"id": target["id"]}, {"$set": {"contact_count": rmsg_count}})
        await finalize_attention(target["id"])
        out = await db.issues.find_one({"id": target["id"]}, {"_id": 0})
        out["matched_existing"] = True
        return out

    # Brand new issue
    issue = Issue(property_id=property_id, resident_id=resident["id"], unit=unit,
                  category=payload.category, description=message, first_reported_at=now_iso(), contact_count=1)
    await db.issues.insert_one(issue.model_dump())
    interaction = Interaction(issue_id=issue.id, resident_id=resident["id"], sender="resident", message=message)
    await db.interactions.insert_one(interaction.model_dump())
    await run_triage(issue.id, interaction.id, message)
    return await db.issues.find_one({"id": issue.id}, {"_id": 0})


@api_router.post("/residents/requests")
async def resident_requests(payload: ResidentLookup):
    prop = await db.properties.find_one({}, {"_id": 0})
    resident = await db.residents.find_one(
        {"name": payload.name, "unit": payload.unit, "property_id": prop["id"] if prop else None},
        {"_id": 0},
    )
    if not resident:
        return {"resident": None, "issues": []}
    issues = await db.issues.find({"resident_id": resident["id"]}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return {"resident": resident, "issues": issues}


@api_router.post("/issues/{issue_id}/reply")
async def resident_reply(issue_id: str, payload: StaffMessage):
    issue = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found")
    interaction = Interaction(issue_id=issue_id, resident_id=issue["resident_id"], sender="resident", message=payload.message)
    await db.interactions.insert_one(interaction.model_dump())

    was_resolved = issue["status"] in ("resolved", "confirmation_pending")
    prior = await db.interactions.find({"issue_id": issue_id}, {"_id": 0}).sort("created_at", 1).to_list(50)
    context = "\n".join(f"[{p['sender']}] {p['message']}" for p in prior[:-1])

    # Analysis-only refresh on a follow-up. NEVER re-route lane/status or auto-resolve an existing ticket.
    analysis = await analyze_message(payload.message, context=context or None)
    order = ["P3", "P2", "P1", "P0"]
    sets = {"contact_count": (issue.get("contact_count") or 1) + 1}
    if analysis:
        await db.interactions.update_one({"id": interaction.id}, {"$set": {
            "detected_intent": analysis.get("primary_intent"),
            "detected_sentiment": analysis.get("sentiment"),
        }})
        if analysis.get("sentiment"):
            sets["sentiment"] = analysis["sentiment"]
        for k in ("safety_risk", "legal_or_financial", "multiple_residents_affected", "interpersonal", "ai_uncertain"):
            if analysis.get(k):
                sets[k] = True
        if analysis.get("human_judgment_required"):
            sets["human_judgment_required"] = True
        new_pri = analysis.get("priority") if analysis.get("priority") in PRIORITY_VALUES else None
        cur_pri = issue.get("priority") or "P2"
        if new_pri and order.index(new_pri) > order.index(cur_pri):  # escalate only, never downgrade
            sets["priority"] = new_pri
            if new_pri == "P0":
                sets.update({"is_emergency": True, "lane": "REVIEW", "human_judgment_required": True})

    escalated_emergency = bool(sets.get("is_emergency")) and not issue.get("is_emergency")

    if was_resolved:
        sets.update({
            "status": "reopened", "lane": "REVIEW", "resident_confirmed": False,
            "failed_resolution": True, "previous_resolved_at": issue.get("resolved_at"),
            "resolution_attempts": (issue.get("resolution_attempts") or 0) + 1,
            "repeat_complaint": await build_repeat_complaint(issue),
        })
    await db.issues.update_one({"id": issue_id}, {"$set": sets})
    if was_resolved or escalated_emergency or sets.get("lane") == "REVIEW":
        await set_review_brief(issue_id, payload.message, context=context or None)
    await finalize_attention(issue_id)
    return interaction


# ------------------- Staff (protected) routes -------------------
async def enrich_issue(issue):
    resident = await db.residents.find_one({"id": issue["resident_id"]}, {"_id": 0})
    issue["resident_name"] = resident["name"] if resident else "Unknown"
    return issue


@api_router.get("/issues")
async def list_issues(user=Depends(get_current_user)):
    issues = await db.issues.find({}, {"_id": 0}).sort("created_at", -1).to_list(2000)
    for i in issues:
        await enrich_issue(i)
    return issues


@api_router.get("/issues/{issue_id}")
async def get_issue(issue_id: str, user=Depends(get_current_user)):
    issue = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found")
    await enrich_issue(issue)
    interactions = await db.interactions.find({"issue_id": issue_id}, {"_id": 0}).sort("created_at", 1).to_list(1000)
    issue["interactions"] = interactions
    return issue


@api_router.patch("/issues/{issue_id}")
async def update_issue(issue_id: str, payload: IssueUpdate, user=Depends(get_current_user)):
    issue = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found")

    updates = {}
    logs = []
    ask_confirmation = False
    if payload.status is not None:
        if payload.status not in STATUS_VALUES:
            raise HTTPException(status_code=400, detail="Invalid status")
        if payload.status == "resolved" and issue.get("lane") == "ACTION" and not issue.get("resident_confirmed"):
            # Step 7: don't resolve immediately — ask the resident to confirm first
            updates["status"] = "confirmation_pending"
            updates["confirmation_requested_at"] = now_iso()
            logs.append("Marked complete — awaiting resident confirmation")
            ask_confirmation = True
        else:
            updates["status"] = payload.status
            logs.append(f"Status changed to {payload.status}")
            if payload.status == "resolved":
                updates["resolved_at"] = now_iso()
    if payload.lane is not None:
        if payload.lane not in LANE_VALUES:
            raise HTTPException(status_code=400, detail="Invalid lane")
        updates["lane"] = payload.lane
        logs.append(f"Lane set to {payload.lane}")
    if payload.priority is not None:
        if payload.priority not in PRIORITY_VALUES:
            raise HTTPException(status_code=400, detail="Invalid priority")
        updates["priority"] = payload.priority
        logs.append(f"Priority set to {payload.priority}")
    if payload.assigned_team is not None:
        if payload.assigned_team and payload.assigned_team not in TEAM_VALUES:
            raise HTTPException(status_code=400, detail="Invalid team")
        updates["assigned_team"] = payload.assigned_team or None
        logs.append(f"Assigned to {payload.assigned_team}" if payload.assigned_team else "Unassigned")
    if payload.human_reason is not None:
        updates["human_reason"] = payload.human_reason

    if updates:
        await db.issues.update_one({"id": issue_id}, {"$set": updates})
    if logs:
        interaction = Interaction(
            issue_id=issue_id,
            sender="system",
            message="; ".join(logs) + f" by {user.get('name', 'Staff')}",
        )
        await db.interactions.insert_one(interaction.model_dump())
    if ask_confirmation:
        await db.interactions.insert_one(Interaction(
            issue_id=issue_id, sender="ai",
            message="Is everything working now? Please confirm in your resident portal — 'Yes, it's resolved' or 'No, I still need help'.",
        ).model_dump())
    await finalize_attention(issue_id)

    updated = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    await enrich_issue(updated)
    return updated


class ConfirmPayload(BaseModel):
    confirmed: bool


@api_router.post("/issues/{issue_id}/confirm")
async def resident_confirm(issue_id: str, payload: ConfirmPayload):
    issue = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found")
    if payload.confirmed:
        await db.issues.update_one({"id": issue_id}, {"$set": {
            "status": "resolved", "resident_confirmed": True,
            "confirmed_at": now_iso(), "resolved_at": issue.get("resolved_at") or now_iso(),
        }})
        await db.interactions.insert_one(Interaction(
            issue_id=issue_id, resident_id=issue["resident_id"], sender="resident",
            message="✓ Confirmed resolved — everything is working now.").model_dump())
    else:
        await db.issues.update_one({"id": issue_id}, {"$set": {
            "status": "reopened", "resident_confirmed": False, "lane": "REVIEW",
            "failed_resolution": True, "previous_resolved_at": issue.get("resolved_at"),
            "resolution_attempts": (issue.get("resolution_attempts") or 0) + 1,
            "human_judgment_required": True,
            "human_reason": "Resident reported the issue is still not resolved after staff marked it complete.",
            "repeat_complaint": await build_repeat_complaint(issue),
        }})
        await db.interactions.insert_one(Interaction(
            issue_id=issue_id, resident_id=issue["resident_id"], sender="resident",
            message="✗ Still needs help — the problem is not resolved.").model_dump())
        await set_review_brief(issue_id, "Resident reports the issue is still not resolved after staff marked it complete.")
    await finalize_attention(issue_id)
    return await db.issues.find_one({"id": issue_id}, {"_id": 0})


@api_router.post("/issues/{issue_id}/message")
async def staff_message(issue_id: str, payload: StaffMessage, user=Depends(get_current_user)):
    issue = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found")
    interaction = Interaction(
        issue_id=issue_id,
        sender="staff",
        message=payload.message,
    )
    await db.interactions.insert_one(interaction.model_dump())
    return interaction


@api_router.get("/stats")
async def stats(user=Depends(get_current_user)):
    all_issues = await db.issues.find({}, {"_id": 0}).to_list(5000)
    total = len(all_issues)
    by_status = {}
    for i in all_issues:
        by_status[i["status"]] = by_status.get(i["status"], 0) + 1
    return {
        "total": total,
        "open": by_status.get("open", 0),
        "in_progress": by_status.get("in_progress", 0),
        "resolved": by_status.get("resolved", 0),
    }


@api_router.get("/dashboard")
async def dashboard(user=Depends(get_current_user)):
    issues = await db.issues.find({}, {"_id": 0}).to_list(5000)
    today = datetime.now(timezone.utc).date().isoformat()
    inter_today = await db.interactions.count_documents({"sender": "resident", "created_at": {"$gte": today}})
    handled = sum(1 for i in issues if i.get("lane") == "RESOLVE")
    actions = sum(1 for i in issues if i.get("lane") == "ACTION")
    reviews = sum(1 for i in issues if i.get("lane") == "REVIEW")
    failed = sum(1 for i in issues if i.get("failed_resolution") or (i.get("resolution_attempts") or 0) > 0)
    conf_pending = sum(1 for i in issues if i.get("status") == "confirmation_pending")
    resolved_issues = [i for i in issues if i.get("status") == "resolved"]
    confirmed = sum(1 for i in resolved_issues if i.get("resident_confirmed"))
    conf_rate = round(100 * confirmed / len(resolved_issues)) if resolved_issues else 0

    frts = []
    for i in issues:
        inters = await db.interactions.find({"issue_id": i["id"]}, {"_id": 0}).sort("created_at", 1).to_list(200)
        first_res = next((x for x in inters if x["sender"] == "resident"), None)
        first_reply = next((x for x in inters if x["sender"] in ("staff", "ai", "system")), None)
        if first_res and first_reply:
            try:
                dt = (datetime.fromisoformat(first_reply["created_at"]) - datetime.fromisoformat(first_res["created_at"])).total_seconds()
                if dt >= 0:
                    frts.append(dt)
            except Exception:
                pass
    median_frt = 0
    if frts:
        frts.sort()
        median_frt = frts[len(frts) // 2]

    return {
        "total": len(issues),
        "resident_interactions_today": inter_today,
        "handled_automatically": handled,
        "actions_created": actions,
        "human_reviews": reviews,
        "failed_resolutions": failed,
        "confirmation_pending": conf_pending,
        "resident_confirmed_rate": conf_rate,
        "median_first_response_seconds": int(median_frt),
    }


@api_router.get("/incidents/detect")
async def detect_incidents(user=Depends(get_current_user)):
    cutoff = (datetime.now(timezone.utc) - timedelta(hours=24)).isoformat()
    issues = await db.issues.find(
        {"created_at": {"$gte": cutoff}, "status": {"$ne": "resolved"},
         "$or": [{"incident_id": None}, {"incident_id": {"$exists": False}}]},
        {"_id": 0},
    ).to_list(1000)

    # Cluster by a SHARED significant keyword within a SHORT window, requiring the SAME category
    from collections import defaultdict
    WINDOW_MIN = 180
    GENERIC = {"apartment", "building", "please", "help", "broken", "working", "need", "unit",
               "door", "room", "thing", "issue", "problem", "today", "morning", "night", "still"}
    kw_map = defaultdict(list)
    for i in issues:
        ents = " ".join(i.get("entities") or [])
        blob = (i.get("description", "") + " " + ents).lower()
        terms = {w for w in re.findall(r"[a-z]+", blob) if w not in _STOPWORDS and w not in GENERIC and len(w) > 3}
        for t in terms:
            kw_map[t].append(i)

    incidents = []
    seen = set()
    for kw, items in sorted(kw_map.items(), key=lambda x: -len(x[1])):
        bycat = defaultdict(list)
        for it in items:
            bycat[it.get("category") or "other"].append(it)
        for cat, catitems in bycat.items():
            residents = {it["resident_id"] for it in catitems}
            if len(residents) < 3:
                continue
            try:
                times = sorted(datetime.fromisoformat(it["created_at"]) for it in catitems)
                window = (times[-1] - times[0]).total_seconds() / 60
            except Exception:
                window = 0
            if window > WINDOW_MIN:
                continue
            ids = frozenset(it["id"] for it in catitems)
            if ids & seen:
                continue
            seen |= ids
            incidents.append({
                "category": cat, "keyword": kw, "count": len(catitems),
                "resident_count": len(residents), "window_minutes": int(window),
                "issue_ids": list(ids), "units": [it["unit"] for it in catitems],
            })
    return {"incidents": incidents}


class MergePayload(BaseModel):
    issue_ids: List[str]
    label: Optional[str] = None


@api_router.post("/incidents/merge")
async def merge_incident(payload: MergePayload, user=Depends(get_current_user)):
    if len(payload.issue_ids) < 2:
        raise HTTPException(status_code=400, detail="Need at least two issues to merge")
    incident_id = f"inc_{uuid.uuid4().hex[:10]}"
    await db.incidents.insert_one({
        "id": incident_id, "label": payload.label or "Shared incident",
        "issue_ids": payload.issue_ids, "created_at": now_iso(), "created_by": user.get("name"),
    })
    await db.issues.update_many({"id": {"$in": payload.issue_ids}}, {"$set": {"incident_id": incident_id}})
    return {"id": incident_id, "issue_ids": payload.issue_ids}


# ------------------- Resident chat thread -------------------
class ThreadLookup(BaseModel):
    name: str
    unit: str
    issue_id: str


@api_router.post("/residents/thread")
async def resident_thread(payload: ThreadLookup):
    prop = await db.properties.find_one({}, {"_id": 0})
    resident = await db.residents.find_one(
        {"name": payload.name, "unit": payload.unit, "property_id": prop["id"] if prop else None}, {"_id": 0}
    )
    if not resident:
        raise HTTPException(status_code=404, detail="Resident not found")
    issue = await db.issues.find_one({"id": payload.issue_id, "resident_id": resident["id"]}, {"_id": 0})
    if not issue:
        raise HTTPException(status_code=404, detail="Request not found")
    interactions = await db.interactions.find({"issue_id": payload.issue_id}, {"_id": 0}).sort("created_at", 1).to_list(1000)
    return {"issue": issue, "interactions": interactions}


# ------------------- Answer approval (staff) -------------------
class ApproveAnswer(BaseModel):
    answer: str


@api_router.post("/issues/{issue_id}/approve-answer")
async def approve_answer(issue_id: str, payload: ApproveAnswer, user=Depends(get_current_user)):
    issue = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found")
    answer = payload.answer.strip()
    if not answer:
        raise HTTPException(status_code=400, detail="Answer cannot be empty")
    if issue.get("lane") != "REVIEW" or not issue.get("suggested_response"):
        raise HTTPException(status_code=400, detail="Answer approval is only available for a REVIEW issue that has a suggested response")
    await db.interactions.insert_one(Interaction(
        issue_id=issue_id, resident_id=issue["resident_id"], sender="ai", message=answer).model_dump())
    await db.interactions.insert_one(Interaction(
        issue_id=issue_id, sender="system", message=f"Suggested answer approved and sent by {user.get('name', 'Staff')}").model_dump())
    await db.issues.update_one({"id": issue_id}, {"$set": {
        "lane": "RESOLVE", "status": "resolved", "resolved_at": now_iso(),
        "auto_response": answer, "answer_confidence": issue.get("answer_confidence") or "high",
        "suggested_response": None, "human_judgment_required": False, "assigned_team": None,
        "policy_conflict": False, "conflicting_documents": [],
        "human_attention_score": 0, "attention_reasons": ["Resolved via approved answer"],
    }})
    updated = await db.issues.find_one({"id": issue_id}, {"_id": 0})
    await enrich_issue(updated)
    return updated


# ------------------- Trend insights (staff) -------------------
@api_router.get("/insights")
async def insights(user=Depends(get_current_user)):
    from collections import Counter
    issues = await db.issues.find({}, {"_id": 0}).to_list(5000)
    repeat = [i for i in issues if (i.get("resolution_attempts") or 0) > 0 or i.get("failed_resolution") or i.get("status") == "reopened" or i.get("repeat_complaint")]
    by_category = Counter((i.get("category") or "other") for i in repeat)
    by_unit = Counter(i.get("unit") for i in repeat)
    weekly = Counter()
    for i in repeat:
        try:
            d = datetime.fromisoformat(i["created_at"])
            iso = d.isocalendar()
            weekly[f"{iso[0]}-W{iso[1]:02d}"] += 1
        except Exception:
            pass
    return {
        "total_repeat": len(repeat),
        "by_category": [{"name": k, "count": v} for k, v in by_category.most_common(10)],
        "by_unit": [{"unit": k, "count": v} for k, v in by_unit.most_common(10)],
        "weekly": [{"week": k, "count": weekly[k]} for k in sorted(weekly.keys())],
    }


@api_router.get("/impact")
async def impact(user=Depends(get_current_user)):
    issues = await db.issues.find({}, {"_id": 0}).to_list(5000)
    total = len(issues)
    automated = sum(1 for i in issues if i.get("lane") in ("RESOLVE", "ACTION"))
    reviews = sum(1 for i in issues if i.get("lane") == "REVIEW")
    repeat = sum(1 for i in issues if i.get("repeat_complaint"))
    failed = sum(1 for i in issues if i.get("failed_resolution") or (i.get("resolution_attempts") or 0) > 0)
    duplicates_prevented = sum(max((i.get("contact_count") or 1) - 1, 0) for i in issues)
    resolved_issues = [i for i in issues if i.get("status") == "resolved"]
    confirmed = sum(1 for i in resolved_issues if i.get("resident_confirmed"))
    conf_rate = round(100 * confirmed / len(resolved_issues)) if resolved_issues else 0
    automation_rate = round(100 * automated / total) if total else 0
    MIN_PER = 8
    hours_saved = round(automated * MIN_PER / 60, 1)
    dash = await dashboard(user)
    return {
        "total_interactions": total,
        "automation_rate": automation_rate,
        "interactions_automated": automated,
        "human_reviews": reviews,
        "median_first_response_seconds": dash["median_first_response_seconds"],
        "resident_confirmed_rate": conf_rate,
        "repeat_complaints": repeat,
        "duplicates_prevented": duplicates_prevented,
        "failed_resolutions": failed,
        "hours_saved": hours_saved,
        "assumed_minutes_per_interaction": MIN_PER,
    }


@api_router.post("/demo/reset")
async def demo_reset(user=Depends(get_current_user)):
    import subprocess
    def _seed():
        return subprocess.run(["python", str(ROOT_DIR / "seed_demo.py")], capture_output=True, text=True, timeout=120)
    r = await run_in_threadpool(_seed)
    ok = r.returncode == 0
    return {"ok": ok, "output": (r.stdout or r.stderr).strip()[-300:]}


@api_router.get("/")
async def root():
    return {"message": "CloseLoop API"}


# ------------------- Property Knowledge (documents) -------------------
async def process_document(doc_id, data, ext):
    await db.property_documents.update_one({"id": doc_id}, {"$set": {"processing_status": "processing"}})
    await db.document_chunks.delete_many({"document_id": doc_id})
    text = await run_in_threadpool(extract_text, data, ext)
    chunks = chunk_text(text)
    if chunks:
        chunk_docs = [{
            "id": str(uuid.uuid4()),
            "document_id": doc_id,
            "chunk_index": idx,
            "text": c,
            "embedding": None,  # reserved for later semantic search
            "created_at": now_iso(),
        } for idx, c in enumerate(chunks)]
        await db.document_chunks.insert_many(chunk_docs)
        await db.property_documents.update_one(
            {"id": doc_id}, {"$set": {"processing_status": "ready", "chunk_count": len(chunks)}}
        )
    else:
        await db.property_documents.update_one(
            {"id": doc_id}, {"$set": {"processing_status": "failed", "chunk_count": 0}}
        )


async def resolve_property(property_id):
    prop = None
    if property_id:
        prop = await db.properties.find_one({"id": property_id}, {"_id": 0})
    if not prop:
        prop = await db.properties.find_one({}, {"_id": 0})
    if not prop:
        raise HTTPException(status_code=400, detail="No property configured")
    return prop


@api_router.get("/documents")
async def list_documents(user=Depends(get_current_user)):
    docs = await db.property_documents.find({"is_deleted": False}, {"_id": 0}).sort("uploaded_at", -1).to_list(1000)
    return docs


@api_router.post("/documents")
async def upload_document(
    file: UploadFile = File(...),
    name: str = Form(...),
    doc_type: str = Form(...),
    property_id: Optional[str] = Form(None),
    user=Depends(get_current_user),
):
    ext = (file.filename.rsplit(".", 1)[-1] if "." in file.filename else "").lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX or TXT.")
    if doc_type not in DOCUMENT_TYPES:
        raise HTTPException(status_code=400, detail="Invalid document type")
    prop = await resolve_property(property_id)
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    content_type = MIME_TYPES.get(ext, file.content_type or "application/octet-stream")
    path = f"{APP_NAME}/{prop['id']}/{uuid.uuid4()}.{ext}"
    result = await run_in_threadpool(put_object, path, data, content_type)
    doc = PropertyDocument(
        property_id=prop["id"],
        name=name.strip() or file.filename,
        doc_type=doc_type,
        storage_path=result["path"],
        original_filename=file.filename,
        content_type=content_type,
        size=result.get("size", len(data)),
        uploaded_by=user.get("name"),
    )
    await db.property_documents.insert_one(doc.model_dump())
    await process_document(doc.id, data, ext)
    return await db.property_documents.find_one({"id": doc.id}, {"_id": 0})


@api_router.put("/documents/{doc_id}/replace")
async def replace_document(doc_id: str, file: UploadFile = File(...), user=Depends(get_current_user)):
    doc = await db.property_documents.find_one({"id": doc_id, "is_deleted": False}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    ext = (file.filename.rsplit(".", 1)[-1] if "." in file.filename else "").lower()
    if ext not in ALLOWED_EXT:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX or TXT.")
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    content_type = MIME_TYPES.get(ext, file.content_type or "application/octet-stream")
    path = f"{APP_NAME}/{doc['property_id']}/{uuid.uuid4()}.{ext}"
    result = await run_in_threadpool(put_object, path, data, content_type)
    await db.property_documents.update_one({"id": doc_id}, {"$set": {
        "storage_path": result["path"],
        "original_filename": file.filename,
        "content_type": content_type,
        "size": result.get("size", len(data)),
        "uploaded_at": now_iso(),
        "uploaded_by": user.get("name"),
        "processing_status": "pending",
        "chunk_count": 0,
    }})
    await process_document(doc_id, data, ext)
    return await db.property_documents.find_one({"id": doc_id}, {"_id": 0})


@api_router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str, user=Depends(get_current_user)):
    res = await db.property_documents.update_one(
        {"id": doc_id, "is_deleted": False}, {"$set": {"is_deleted": True}}
    )
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    await db.document_chunks.delete_many({"document_id": doc_id})
    return {"ok": True}


@api_router.get("/documents/{doc_id}/download")
async def download_document(doc_id: str, user=Depends(get_current_user)):
    doc = await db.property_documents.find_one({"id": doc_id, "is_deleted": False}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    data, content_type = await run_in_threadpool(get_object, doc["storage_path"])
    return Response(
        content=data,
        media_type=doc.get("content_type", content_type),
        headers={"Content-Disposition": f'inline; filename="{doc["original_filename"]}"'},
    )


app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()


@app.on_event("startup")
async def startup_storage():
    try:
        init_storage()
        logger.info("Object storage initialized")
    except Exception as e:
        logger.error(f"Storage init failed: {e}")
